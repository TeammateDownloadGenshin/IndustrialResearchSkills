#!/usr/bin/env python3
"""Render a PDF or DOCX report and extract layout-aware preparation metadata."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import zipfile
from datetime import date, datetime
from pathlib import Path
from typing import Any, Iterable

import pdfplumber
from PIL import Image, ImageDraw, ImageOps
from pypdf import PdfReader

try:
    from docx import Document
except ImportError:  # pragma: no cover - reported through the CLI
    Document = None


SCHEMA_VERSION = "1.0"
TEXT_DATE_PATTERNS = (
    re.compile(r"\b((?:19|20)\d{2})[-/.](0?[1-9]|1[0-2])[-/.](0?[1-9]|[12]\d|3[01])\b"),
    re.compile(r"\b((?:19|20)\d{2})(0[1-9]|1[0-2])(0[1-9]|[12]\d|3[01])\b"),
    re.compile(
        r"\b(January|February|March|April|May|June|July|August|September|October|November|December)"
        r"\s+(\d{1,2})(?:st|nd|rd|th)?[,]?\s+((?:19|20)\d{2})\b",
        re.IGNORECASE,
    ),
    re.compile(
        r"\b((?:19|20)\d{2})\s*\u5e74\s*(0?[1-9]|1[0-2])\s*\u6708\s*"
        r"(0?[1-9]|[12]\d|3[01])\s*\u65e5"
    ),
)
FILENAME_DATE_PATTERN = re.compile(r"(?<!\d)((?:19|20)\d{2})(0[1-9]|1[0-2])(0[1-9]|[12]\d|3[01])(?!\d)")
PDF_DATE_PATTERN = re.compile(r"D:((?:19|20)\d{2})(\d{2})(\d{2})")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Render a PDF or DOCX report and write report.json with page images and extracted structure."
    )
    parser.add_argument("input", type=Path, help="Input PDF or DOCX report.")
    parser.add_argument("--work-dir", required=True, type=Path, help="Writable run directory.")
    parser.add_argument("--dpi", type=int, default=220, help="Page render resolution. Default: 220.")
    parser.add_argument("--pdftoppm", type=Path, help="Explicit pdftoppm executable path.")
    parser.add_argument("--soffice", type=Path, help="Explicit LibreOffice soffice executable path.")
    parser.add_argument(
        "--report-date",
        help="Visually confirmed cover or body date in YYYY-MM-DD format, for example on an image-only page.",
    )
    parser.add_argument("--output", type=Path, help="Output JSON path. Default: <work-dir>/report.json.")
    return parser.parse_args()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def find_executable(explicit: Path | None, env_name: str, candidates: Iterable[str]) -> str:
    if explicit:
        resolved = explicit.expanduser().resolve()
        if not resolved.is_file():
            raise FileNotFoundError(f"Executable does not exist: {resolved}")
        return str(resolved)
    env_value = os.environ.get(env_name)
    if env_value:
        resolved = Path(env_value).expanduser().resolve()
        if not resolved.is_file():
            raise FileNotFoundError(f"{env_name} does not point to a file: {resolved}")
        return str(resolved)
    for candidate in candidates:
        found = shutil.which(candidate)
        if found:
            return found
    raise FileNotFoundError(
        f"Required executable was not found. Provide --{env_name.lower().replace('_path', '').replace('_', '-')} "
        f"or set {env_name}."
    )


def valid_date(year: int, month: int, day: int) -> date | None:
    try:
        return date(year, month, day)
    except ValueError:
        return None


def find_text_date(text: str) -> date | None:
    month_names = {
        name.lower(): index
        for index, name in enumerate(
            (
                "January",
                "February",
                "March",
                "April",
                "May",
                "June",
                "July",
                "August",
                "September",
                "October",
                "November",
                "December",
            ),
            start=1,
        )
    }
    for pattern_index, pattern in enumerate(TEXT_DATE_PATTERNS):
        match = pattern.search(text)
        if not match:
            continue
        if pattern_index == 2:
            month = month_names[match.group(1).lower()]
            candidate = valid_date(int(match.group(3)), month, int(match.group(2)))
        else:
            candidate = valid_date(int(match.group(1)), int(match.group(2)), int(match.group(3)))
        if candidate:
            return candidate
    return None


def parse_metadata_date(value: Any) -> date | None:
    if value is None:
        return None
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    text = str(value)
    pdf_match = PDF_DATE_PATTERN.search(text)
    if pdf_match:
        return valid_date(int(pdf_match.group(1)), int(pdf_match.group(2)), int(pdf_match.group(3)))
    return find_text_date(text)


def filename_date(path: Path) -> date | None:
    match = FILENAME_DATE_PATTERN.search(path.stem)
    if not match:
        return None
    return valid_date(int(match.group(1)), int(match.group(2)), int(match.group(3)))


def clean_output_images(directory: Path) -> None:
    directory.mkdir(parents=True, exist_ok=True)
    for pattern in ("raw-*.png", "page-*.png", "contact-*.png"):
        for path in directory.glob(pattern):
            if path.is_file():
                path.unlink()


def run_command(command: list[str], timeout: int, label: str) -> subprocess.CompletedProcess[str]:
    result = subprocess.run(
        command,
        check=False,
        capture_output=True,
        text=True,
        timeout=timeout,
    )
    if result.returncode != 0:
        raise RuntimeError(
            f"{label} failed with exit code {result.returncode}. "
            f"stdout={result.stdout.strip()} stderr={result.stderr.strip()}"
        )
    return result


def convert_docx_to_pdf(input_path: Path, work_dir: Path, soffice: str) -> Path:
    converted_dir = work_dir / "converted"
    profile_dir = work_dir / "libreoffice-profile"
    converted_dir.mkdir(parents=True, exist_ok=True)
    profile_dir.mkdir(parents=True, exist_ok=True)
    expected = converted_dir / f"{input_path.stem}.pdf"
    if expected.exists():
        expected.unlink()
    profile_uri = profile_dir.resolve().as_uri()
    command = [
        soffice,
        "--headless",
        f"-env:UserInstallation={profile_uri}",
        "--convert-to",
        "pdf",
        "--outdir",
        str(converted_dir),
        str(input_path),
    ]
    run_command(command, timeout=240, label="DOCX to PDF conversion")
    if not expected.is_file() or expected.stat().st_size == 0:
        candidates = sorted(converted_dir.glob("*.pdf"))
        if len(candidates) != 1:
            raise RuntimeError("LibreOffice completed but did not produce the expected PDF.")
        expected = candidates[0]
    rendered_pdf = work_dir / "rendered-report.pdf"
    shutil.copy2(expected, rendered_pdf)
    return rendered_pdf


def extract_docx_structure(path: Path) -> tuple[dict[str, Any], str | None, date | None]:
    if Document is None:
        raise RuntimeError("python-docx is unavailable in the selected Python runtime.")
    document = Document(str(path))
    paragraphs: list[dict[str, Any]] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        paragraphs.append(
            {
                "id": f"P{index:05d}",
                "order": index,
                "style": paragraph.style.name if paragraph.style else None,
                "text": text,
            }
        )
    tables: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"id": f"T{table_index:04d}", "order": table_index, "rows": rows})
    core = document.core_properties
    media_count = 0
    with zipfile.ZipFile(path) as archive:
        media_count = sum(1 for name in archive.namelist() if name.startswith("word/media/") and not name.endswith("/"))
    metadata_date = parse_metadata_date(core.modified) or parse_metadata_date(core.created)
    title = (core.title or "").strip() or None
    return (
        {
            "paragraphs": paragraphs,
            "tables": tables,
            "embedded_media_count": media_count,
            "core_properties": {
                "title": title,
                "subject": (core.subject or "").strip() or None,
                "author": (core.author or "").strip() or None,
                "created": core.created.isoformat() if core.created else None,
                "modified": core.modified.isoformat() if core.modified else None,
            },
        },
        title,
        metadata_date,
    )


def render_pdf(pdf_path: Path, pages_dir: Path, pdftoppm: str, dpi: int) -> list[Path]:
    clean_output_images(pages_dir)
    prefix = pages_dir / "raw"
    run_command(
        [pdftoppm, "-png", "-r", str(dpi), str(pdf_path), str(prefix)],
        timeout=600,
        label="PDF rendering",
    )
    raw_pages = sorted(
        pages_dir.glob("raw-*.png"),
        key=lambda item: int(re.search(r"(\d+)$", item.stem).group(1)),
    )
    if not raw_pages:
        raise RuntimeError("PDF rendering produced no page images.")
    output_pages: list[Path] = []
    for page_number, raw_path in enumerate(raw_pages, start=1):
        output_path = pages_dir / f"page-{page_number:04d}.png"
        raw_path.replace(output_path)
        output_pages.append(output_path.resolve())
    return output_pages


def extract_pdf_pages(pdf_path: Path, page_images: list[Path]) -> tuple[list[dict[str, Any]], str | None, date | None]:
    reader = PdfReader(str(pdf_path))
    metadata = reader.metadata
    metadata_title = None
    metadata_date = None
    if metadata:
        metadata_title = (getattr(metadata, "title", None) or "").strip() or None
        metadata_date = parse_metadata_date(getattr(metadata, "creation_date", None)) or parse_metadata_date(
            getattr(metadata, "modification_date", None)
        )
    pages: list[dict[str, Any]] = []
    with pdfplumber.open(str(pdf_path)) as document:
        if len(document.pages) != len(page_images):
            raise RuntimeError(
                f"Rendered page count {len(page_images)} does not match PDF page count {len(document.pages)}."
            )
        for page_number, (page, image_path) in enumerate(zip(document.pages, page_images), start=1):
            try:
                text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
            except Exception:
                text = ""
            words: list[dict[str, Any]] = []
            try:
                for word in page.extract_words(use_text_flow=True, keep_blank_chars=False):
                    words.append(
                        {
                            "text": word.get("text", ""),
                            "x0": round(float(word.get("x0", 0.0)), 3),
                            "top": round(float(word.get("top", 0.0)), 3),
                            "x1": round(float(word.get("x1", 0.0)), 3),
                            "bottom": round(float(word.get("bottom", 0.0)), 3),
                        }
                    )
            except Exception:
                words = []
            with Image.open(image_path) as image:
                width_pixels, height_pixels = image.size
            compact_length = len(re.sub(r"\s+", "", text))
            pages.append(
                {
                    "page": page_number,
                    "image_path": str(image_path),
                    "width_points": round(float(page.width), 3),
                    "height_points": round(float(page.height), 3),
                    "width_pixels": width_pixels,
                    "height_pixels": height_pixels,
                    "text": text,
                    "text_length": compact_length,
                    "vision_required": compact_length < 20,
                    "words": words,
                }
            )
    return pages, metadata_title, metadata_date


def create_contact_sheets(page_images: list[Path], output_dir: Path) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    for existing in output_dir.glob("contact-*.png"):
        if existing.is_file():
            existing.unlink()
    results: list[Path] = []
    page_width = 360
    page_height = 480
    margin = 24
    label_height = 28
    columns = 2
    rows = 3
    per_sheet = columns * rows
    for sheet_index in range(0, len(page_images), per_sheet):
        subset = page_images[sheet_index : sheet_index + per_sheet]
        canvas = Image.new(
            "RGB",
            (
                margin + columns * (page_width + margin),
                margin + rows * (page_height + label_height + margin),
            ),
            "white",
        )
        draw = ImageDraw.Draw(canvas)
        for local_index, image_path in enumerate(subset):
            column = local_index % columns
            row = local_index // columns
            left = margin + column * (page_width + margin)
            top = margin + row * (page_height + label_height + margin)
            with Image.open(image_path) as source:
                thumbnail = ImageOps.contain(source.convert("RGB"), (page_width, page_height))
            image_left = left + (page_width - thumbnail.width) // 2
            image_top = top + (page_height - thumbnail.height) // 2
            canvas.paste(thumbnail, (image_left, image_top))
            page_number = sheet_index + local_index + 1
            draw.rectangle((left, top, left + page_width, top + page_height), outline="#B7B7B7", width=1)
            draw.text((left, top + page_height + 6), f"Page {page_number}", fill="#222222")
        output_path = output_dir / f"contact-{sheet_index // per_sheet + 1:03d}.png"
        canvas.save(output_path, format="PNG", optimize=True)
        results.append(output_path.resolve())
    return results


def choose_report_date(
    input_path: Path,
    first_page_text: str,
    docx_structure: dict[str, Any] | None,
    metadata_date: date | None,
) -> tuple[date, str, bool]:
    document_text_parts = [first_page_text]
    if docx_structure:
        document_text_parts.extend(item["text"] for item in docx_structure.get("paragraphs", [])[:30])
    explicit_date = find_text_date("\n".join(document_text_parts))
    if explicit_date:
        return explicit_date, "document_text", False
    if metadata_date:
        return metadata_date, "document_metadata", False
    name_date = filename_date(input_path)
    if name_date:
        return name_date, "filename", False
    modified_date = datetime.fromtimestamp(input_path.stat().st_mtime).date()
    return modified_date, "file_modified_time", True


def main() -> int:
    args = parse_args()
    input_path = args.input.expanduser().resolve()
    work_dir = args.work_dir.expanduser().resolve()
    output_path = (args.output or work_dir / "report.json").expanduser().resolve()
    if not input_path.is_file():
        raise FileNotFoundError(f"Input report does not exist: {input_path}")
    input_type = input_path.suffix.lower().lstrip(".")
    if input_type not in {"pdf", "docx"}:
        raise ValueError("Input must be a PDF or DOCX file.")
    if args.dpi < 100 or args.dpi > 400:
        raise ValueError("DPI must be between 100 and 400.")
    work_dir.mkdir(parents=True, exist_ok=True)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    pdftoppm = find_executable(args.pdftoppm, "PDFTOPPM_PATH", ("pdftoppm", "pdftoppm.exe"))
    docx_structure = None
    docx_title = None
    docx_metadata_date = None
    if input_type == "docx":
        docx_structure, docx_title, docx_metadata_date = extract_docx_structure(input_path)
        soffice = find_executable(args.soffice, "SOFFICE_PATH", ("soffice", "soffice.com", "soffice.exe"))
        pdf_path = convert_docx_to_pdf(input_path, work_dir, soffice)
    else:
        pdf_path = input_path

    page_images = render_pdf(pdf_path, work_dir / "pages", pdftoppm, args.dpi)
    pages, pdf_title, pdf_metadata_date = extract_pdf_pages(pdf_path, page_images)
    metadata_date = docx_metadata_date or pdf_metadata_date
    report_date, date_source, date_is_fallback = choose_report_date(
        input_path,
        "\n".join(page["text"] for page in pages[:3]),
        docx_structure,
        metadata_date,
    )
    if args.report_date:
        try:
            report_date = date.fromisoformat(args.report_date)
        except ValueError as error:
            raise ValueError("--report-date must use YYYY-MM-DD format.") from error
        date_source = "document_visual_text"
        date_is_fallback = False
    contact_sheets = create_contact_sheets(page_images, work_dir / "contact-sheets")
    title = docx_title or pdf_title or input_path.stem
    payload = {
        "schema_version": SCHEMA_VERSION,
        "report": {
            "input_path": str(input_path),
            "input_filename": input_path.name,
            "input_type": input_type,
            "title": title,
            "report_date": report_date.isoformat(),
            "report_date_source": date_source,
            "report_date_is_fallback": date_is_fallback,
            "sha256": sha256_file(input_path),
            "page_count": len(pages),
            "rendered_pdf_path": str(pdf_path.resolve()),
            "render_dpi": args.dpi,
        },
        "pages": pages,
        "docx_structure": docx_structure,
        "contact_sheets": [str(path) for path in contact_sheets],
    }
    output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(
        json.dumps(
            {
                "status": "ok",
                "output": str(output_path),
                "page_count": len(pages),
                "report_date": report_date.isoformat(),
                "report_date_source": date_source,
                "vision_required_pages": [page["page"] for page in pages if page["vision_required"]],
            },
            ensure_ascii=True,
        )
    )
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(json.dumps({"status": "error", "error": str(exc)}, ensure_ascii=True), file=sys.stderr)
        raise SystemExit(1)
